import pdfplumber
import docx
import os
import io

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse PDF document: {str(e)}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Detects file extension and extracts all plain text in-memory from bytes."""
    _, ext = os.path.splitext(filename.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF, DOCX, or TXT file.")
